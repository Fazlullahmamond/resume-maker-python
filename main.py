import sys
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QPushButton,
    QFileDialog, QListWidget, QLineEdit, QHBoxLayout, QMessageBox, QInputDialog
)
from PyQt6.QtCore import Qt
from docx import Document
from fpdf import FPDF
import re
from datetime import datetime

# ------------------------
# Validation Functions
# ------------------------
def validate_email(email):
    pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    return re.match(pattern, email) is not None

def validate_phone(phone):
    pattern = r'^\+?\d{7,15}$'
    return re.match(pattern, phone) is not None

def validate_date(date_text):
    try:
        datetime.strptime(date_text, "%Y-%m-%d")
        return True
    except ValueError:
        return False

# ------------------------
# Resume Generator
# ------------------------
def generate_docx(resume_data, filename="resume.docx"):
    doc = Document()
    doc.add_heading(resume_data['name'], 0)
    doc.add_paragraph(f"Email: {resume_data['email']}")
    doc.add_paragraph(f"Phone: {resume_data['phone']}")
    doc.add_paragraph(f"Address: {resume_data['address']}")
    
    doc.add_heading("Education", level=1)
    for edu in resume_data['education']:
        doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['start']} to {edu['end']})")
    
    doc.add_heading("Experience", level=1)
    for exp in resume_data['experience']:
        doc.add_paragraph(f"{exp['title']} - {exp['company']} ({exp['start']} to {exp['end']})")
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(resume_data['skills']))
    
    doc.save(filename)

def generate_pdf(resume_data, filename="resume.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, resume_data['name'], ln=True)
    
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 8, f"Email: {resume_data['email']}", ln=True)
    pdf.cell(0, 8, f"Phone: {resume_data['phone']}", ln=True)
    pdf.cell(0, 8, f"Address: {resume_data['address']}", ln=True)
    
    pdf.ln(5)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Education", ln=True)
    pdf.set_font("Arial", "", 12)
    for edu in resume_data['education']:
        pdf.cell(0, 8, f"{edu['degree']} - {edu['institution']} ({edu['start']} to {edu['end']})", ln=True)
    
    pdf.ln(3)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Experience", ln=True)
    pdf.set_font("Arial", "", 12)
    for exp in resume_data['experience']:
        pdf.cell(0, 8, f"{exp['title']} - {exp['company']} ({exp['start']} to {exp['end']})", ln=True)
    
    pdf.ln(3)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Skills", ln=True)
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 8, ", ".join(resume_data['skills']))
    
    pdf.output(filename)

# ------------------------
# PyQt6 GUI
# ------------------------
class ResumeMaker(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Resume Maker")
        self.setGeometry(100, 100, 800, 600)

        self.education_entries = []
        self.experience_entries = []

        layout = QVBoxLayout()

        # Personal Info
        layout.addWidget(QLabel("Full Name:"))
        self.name_input = QLineEdit()
        layout.addWidget(self.name_input)

        layout.addWidget(QLabel("Email:"))
        self.email_input = QLineEdit()
        layout.addWidget(self.email_input)

        layout.addWidget(QLabel("Phone:"))
        self.phone_input = QLineEdit()
        layout.addWidget(self.phone_input)

        layout.addWidget(QLabel("Address:"))
        self.address_input = QLineEdit()
        layout.addWidget(self.address_input)

        layout.addWidget(QLabel("Skills (comma separated):"))
        self.skills_input = QLineEdit()
        layout.addWidget(self.skills_input)

        # Buttons for adding education and experience
        self.add_edu_btn = QPushButton("Add Education")
        self.add_edu_btn.clicked.connect(self.add_education_entry)
        layout.addWidget(self.add_edu_btn)

        self.add_exp_btn = QPushButton("Add Experience")
        self.add_exp_btn.clicked.connect(self.add_experience_entry)
        layout.addWidget(self.add_exp_btn)

        # Lists to show added education and experience
        self.edu_list = QListWidget()
        layout.addWidget(QLabel("Education Entries:"))
        layout.addWidget(self.edu_list)

        self.exp_list = QListWidget()
        layout.addWidget(QLabel("Experience Entries:"))
        layout.addWidget(self.exp_list)

        # Output type
        self.pdf_btn = QPushButton("Generate PDF")
        self.pdf_btn.clicked.connect(lambda: self.generate_resume("PDF"))
        self.docx_btn = QPushButton("Generate DOCX")
        self.docx_btn.clicked.connect(lambda: self.generate_resume("DOCX"))
        layout.addWidget(self.pdf_btn)
        layout.addWidget(self.docx_btn)

        self.setLayout(layout)

    def add_education_entry(self):
        degree, ok1 = QInputDialog.getText(self, "Education Entry", "Degree:")
        if not ok1 or not degree.strip(): return
        institution, ok2 = QInputDialog.getText(self, "Education Entry", "Institution:")
        if not ok2 or not institution.strip(): return
        start, ok3 = QInputDialog.getText(self, "Education Entry", "Start Date (YYYY-MM-DD):")
        if not ok3 or not validate_date(start): return
        end, ok4 = QInputDialog.getText(self, "Education Entry", "End Date (YYYY-MM-DD):")
        if not ok4 or not validate_date(end): return
        entry = {"degree": degree, "institution": institution, "start": start, "end": end}
        self.education_entries.append(entry)
        self.edu_list.addItem(f"{degree} - {institution} ({start} to {end})")

    def add_experience_entry(self):
        title, ok1 = QInputDialog.getText(self, "Experience Entry", "Job Title:")
        if not ok1 or not title.strip(): return
        company, ok2 = QInputDialog.getText(self, "Experience Entry", "Company:")
        if not ok2 or not company.strip(): return
        start, ok3 = QInputDialog.getText(self, "Experience Entry", "Start Date (YYYY-MM-DD):")
        if not ok3 or not validate_date(start): return
        end, ok4 = QInputDialog.getText(self, "Experience Entry", "End Date (YYYY-MM-DD or Present):")
        if not ok4 or (end.lower() != "present" and not validate_date(end)): return
        entry = {"title": title, "company": company, "start": start, "end": end}
        self.experience_entries.append(entry)
        self.exp_list.addItem(f"{title} - {company} ({start} to {end})")

    def generate_resume(self, file_type):
        name = self.name_input.text()
        email = self.email_input.text()
        phone = self.phone_input.text()
        address = self.address_input.text()
        skills = [s.strip() for s in self.skills_input.text().split(",")]

        if not name or not validate_email(email) or not validate_phone(phone):
            QMessageBox.critical(self, "Error", "Please provide valid name, email, and phone.")
            return

        resume_data = {
            "name": name,
            "email": email,
            "phone": phone,
            "address": address,
            "skills": skills,
            "education": self.education_entries,
            "experience": self.experience_entries
        }

        filename, _ = QFileDialog.getSaveFileName(
            self,
            f"Save Resume as {file_type}",
            "",
            f"{file_type} Files (*.{file_type.lower()})"
        )
        if not filename: return
        if not filename.lower().endswith(f".{file_type.lower()}"):
            filename += f".{file_type.lower()}"

        if file_type == "PDF":
            generate_pdf(resume_data, filename)
        else:
            generate_docx(resume_data, filename)

        QMessageBox.information(self, "Success", f"Resume saved as {filename}")

# ------------------------
# Run App
# ------------------------
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ResumeMaker()
    window.show()
    sys.exit(app.exec())
